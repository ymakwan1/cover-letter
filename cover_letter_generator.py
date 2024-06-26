import os
import sys
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from docx2pdf import convert

# Load environment variables from .env file
load_dotenv()

def generate_cover_letter(position: str, company: str) -> None:
    template_path = os.getenv('COVER_LETTER_TEMPLATE_PATH')
    font_size_str = os.getenv('COVER_LETTER_FONT_SIZE')
    font_name = os.getenv('COVER_LETTER_FONT_NAME')
    output_dir = os.getenv('COVER_LETTER_OUTPUT_DIR')

    if not template_path:
        print("Cover letter template path not found. Please set the COVER_LETTER_TEMPLATE_PATH environment variable.")
        return

    if not output_dir:
        print("Output directory path not found. Please set the COVER_LETTER_OUTPUT_DIR environment variable.")
        return

    # Convert font size to integer
    font_size = int(font_size_str) if font_size_str else None

    # Load the cover letter template
    doc = Document(template_path)

    # Find and replace placeholders with actual values
    for paragraph in doc.paragraphs:
        if '[POSITION]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[POSITION]', position)
        if '[COMPANY NAME]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[COMPANY NAME]', company)

    # Apply font settings
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if font_size:
                run.font.size = Pt(font_size)
            if font_name:
                run.font.name = font_name

    # Define the output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the modified document as DOCX
    doc_path = os.path.join(output_dir, f"cover_letter_{company}_{position}.docx")
    doc.save(doc_path)

    # Convert the Word document to PDF
    pdf_path = os.path.join(output_dir, f"Yash_Makwana_Cover_Letter_{company}_{position}.pdf")
    convert(doc_path, pdf_path)

    print(f"Cover letter generated and saved as: {pdf_path}")

    # Delete the DOCX file
    os.remove(doc_path)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python3 cover_letter_generator.py <position> <company>")
        sys.exit(1)

    position = sys.argv[1]
    company = sys.argv[2]
    generate_cover_letter(position, company)
